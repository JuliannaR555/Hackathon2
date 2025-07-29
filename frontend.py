import os
import sys
import tempfile
import time

import streamlit as st
from sentence_transformers import SentenceTransformer

# Ajout du dossier aux imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Imports internes
from ingestion.extractor import extract_text
from ingestion.cleaner import clean_text
from ingestion.processing.splitter import split_by_page, score_longueur, couleur_score
from ingestion.processing.embedder import generate_embeddings
from ingestion.processing.indexer import create_faiss_index, index_chunks, search_index
from ingestion.processing.summarizer import summarize_text
from ingestion.processing.export import export_summary_txt, export_summary_pdf

# ============================================
# 📌 CONFIGURATION STREAMLIT
# ============================================
st.set_page_config(page_title="Recherche IA + Résumé", layout="wide")
st.title("🔎 Analyse intelligente de vos documents et archives")

# Initialisation de session
if "doc_chunks" not in st.session_state:
    st.session_state.doc_chunks = []
if "faiss_index" not in st.session_state:
    st.session_state.faiss_index = None
if "embeddings" not in st.session_state:
    st.session_state.embeddings = None

# Modèle SentenceTransformer
model = SentenceTransformer("all-MiniLM-L6-v2")

# ============================================
# 📅 TÉLÉVERSEMENT DE DOCUMENT
# ============================================
st.subheader("📂 1. Charger un fichier (PDF, DOCX, TXT, ODT, HTML)")
uploaded_file = st.file_uploader("Types acceptés :", type=["pdf", "docx", "txt", "odt", "html", "htm"])

if uploaded_file:
    st.info("📄 Traitement du fichier en cours...")
    with open("data/outputs/log.txt", "a", encoding="utf-8") as log:
        log.write(f"[UPLOAD] Fichier reçu : {uploaded_file.name}\n")

    ext = os.path.splitext(uploaded_file.name)[-1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_path = tmp_file.name

    pages = split_by_page(tmp_path)
    with open("data/outputs/log.txt", "a", encoding="utf-8") as log:
        log.write(f"[PARSE] {len(pages)} pages extraites depuis le fichier {uploaded_file.name}\n")
    st.session_state.doc_chunks = pages

    embeddings = generate_embeddings(pages)
    with open("data/outputs/log.txt", "a", encoding="utf-8") as log:
        log.write(f"[EMBED] Embeddings générés ({embeddings.shape})\n")
    st.session_state.embeddings = embeddings

    faiss_index = create_faiss_index(embeddings.shape[1])
    index_chunks(faiss_index, embeddings)
    st.session_state.faiss_index = faiss_index

    st.success(f"✅ {len(pages)} pages ou blocs indexés.")
    os.remove(tmp_path)

# ============================================
# 🔍 INTERFACE DE RECHERCHE
# ============================================
if st.session_state.faiss_index:
    st.subheader("🔍 2. Posez une question sur le document")

    user_query = st.text_input("Votre question :")

    if user_query:
        query_vec = model.encode([user_query], convert_to_numpy=True)
        indices, scores = search_index(st.session_state.faiss_index, query_vec, top_k=3)

        with open("data/outputs/log.txt", "a", encoding="utf-8") as log:
            log.write(f"[QUERY] Question posée : {user_query} | Top-3 résultats : {indices}\n")

        sorted_pairs = sorted(zip(indices, scores), key=lambda x: -x[1])
        sorted_indices = [i for i, _ in sorted_pairs]
        selected_chunks = [st.session_state.doc_chunks[i] for i in sorted_indices]

        st.markdown("### 🧠 Pages les plus pertinentes")
        st.markdown(
            """
Les pages ci-dessous sont celles qui correspondent **le mieux à votre question**, sur le plan **sémantique**.  
Elles sont affichées **dans leur intégralité**, sans découpe artificielle.

<span style='font-size: 0.9em;'>
⬜ Vert : contenu riche et développé  
🟨 Jaune : contenu modéré  
🔴 Rouge : contenu court ou peu informatif
</span>
""",
            unsafe_allow_html=True
        )

        for i, chunk in enumerate(selected_chunks):
            score = score_longueur(chunk)
            bg = couleur_score(score)
            with st.expander(f"📄 Passage {i+1}"):
                st.markdown(f"<div style='background-color:{bg}; padding:10px'>{chunk}</div>", unsafe_allow_html=True)

        full_text = " ".join(selected_chunks)

        # Mesure du temps de calcul
        start_time = time.time()
        summary = summarize_text(full_text)
        with open("data/outputs/log.txt", "a", encoding="utf-8") as log:
            log.write(f"[SUMMARY] Résumé généré ({len(summary.split())} mots)\n")
        duration = time.time() - start_time

        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown("### 📜 Résumé automatique")
        with col2:
            st.markdown(f"\n\n⏱️ *{duration:.2f} sec*")

        st.success(summary)


# ============================================
# 📅 EXPORT DU RÉSUMÉ
# ============================================
if st.button("📅 Télécharger le résumé"):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as txt_file:
        export_summary_txt(summary, txt_file.name)
        st.download_button("⬇️ Fichier .txt", data=open(txt_file.name, "rb"), file_name="resume.txt")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_file:
        export_summary_pdf(summary, pdf_file.name)
        st.download_button("⬇️ Fichier .pdf", data=open(pdf_file.name, "rb"), file_name="resume.pdf")

    from docx import Document
    doc = Document()
    doc.add_heading("Résumé automatique", level=1)
    doc.add_paragraph(summary)
    docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(docx_path)
    st.download_button("⬇️ Fichier .docx", data=open(docx_path, "rb"), file_name="resume.docx")


# ============================================
# 📜 LOGS
# ============================================
st.sidebar.markdown("📜 **Logs**")
if st.sidebar.button("📖 Voir log.txt"):
    log_path = os.path.join("data", "outputs", "log.txt")
    if os.path.exists(log_path):
        with open(log_path, "r", encoding="utf-8") as log_file:
            logs = log_file.read()
        st.sidebar.text_area("Journal :", logs, height=300)
    else:
        st.sidebar.warning("Aucun log disponible.")
